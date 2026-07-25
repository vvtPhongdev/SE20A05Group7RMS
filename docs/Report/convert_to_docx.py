import os
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Set path and import weekly data
sys.path.append(os.path.dirname(os.path.abspath(__file__)))
from generate_ai_report import WEEKLY_DATA

def create_weekly_docx(week_name, tasks):
    doc = docx.Document()
    
    # Set page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Configure default style (Normal)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # The title and introduction paragraphs are removed as requested.
    pass
    
    for idx, task in enumerate(tasks):
        # Heading for each task (Times New Roman, Bold, 12pt, Left aligned)
        heading = doc.add_paragraph()
        heading_run = heading.add_run(f"Nhiệm vụ {idx + 1}: {task['Task / Activity']}")
        heading_run.font.name = 'Times New Roman'
        heading_run.font.size = Pt(12)
        heading_run.font.bold = True
        heading_run.font.color.rgb = RGBColor(0, 0, 0)
        heading.paragraph_format.space_before = Pt(18)
        heading.paragraph_format.space_after = Pt(6)
        heading.paragraph_format.keep_with_next = True
        
        # Metadata Table (All cells must use Times New Roman, 12pt)
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        
        metadata = [
            ("Tên công cụ AI", task["AI Tool Used"]),
            ("Giai đoạn SDLC", task["SDLC Phase"]),
            ("Thước đo định lượng", task["Quantitative Measure"]),
            ("Giá trị đóng góp (1-5)", str(task["Value Added (1-5)"])),
            ("Rủi ro & Hạn chế", task["Risks / Limitations Observed"])
        ]
        
        for r_idx, (key, val) in enumerate(metadata):
            row = table.rows[r_idx]
            
            cell_key = row.cells[0]
            cell_key.text = key
            run_k = cell_key.paragraphs[0].runs[0]
            run_k.font.name = 'Times New Roman'
            run_k.font.size = Pt(12)
            run_k.font.bold = True
            
            cell_val = row.cells[1]
            cell_val.text = val
            run_v = cell_val.paragraphs[0].runs[0]
            run_v.font.name = 'Times New Roman'
            run_v.font.size = Pt(12)
            
        doc.add_paragraph().paragraph_format.space_after = Pt(6)
        
        # 1. Prompt (Times New Roman, Bold, 12pt)
        p_head = doc.add_paragraph()
        p_head_run = p_head.add_run("1. Câu hỏi gửi AI (Prompt)")
        p_head_run.font.name = 'Times New Roman'
        p_head_run.font.size = Pt(12)
        p_head_run.font.bold = True
        p_head.paragraph_format.keep_with_next = True
        
        p_box = doc.add_paragraph()
        p_box.paragraph_format.left_indent = Inches(0.25)
        p_box.paragraph_format.right_indent = Inches(0.25)
        p_box_run = p_box.add_run(task["Prompt"])
        p_box_run.font.name = 'Times New Roman'
        p_box_run.font.size = Pt(12)
        p_box_run.font.italic = True
        p_box.paragraph_format.space_after = Pt(12)
        
        # 2. Response (Times New Roman, Bold, 12pt)
        r_head = doc.add_paragraph()
        r_head_run = r_head.add_run("2. Câu trả lời chi tiết của AI")
        r_head_run.font.name = 'Times New Roman'
        r_head_run.font.size = Pt(12)
        r_head_run.font.bold = True
        r_head.paragraph_format.keep_with_next = True
        
        # Split code blocks
        parts = task["Response"].split("```")
        for p_idx, part in enumerate(parts):
            part_str = part.strip()
            if not part_str:
                continue
                
            if p_idx % 2 == 1:
                # Code block (Times New Roman, 12pt)
                code_lines = part_str.split("\n")
                if code_lines and code_lines[0] in ["typescript", "javascript", "sql", "prisma", "text", "python", "html"]:
                    code_content = "\n".join(code_lines[1:])
                else:
                    code_content = part_str
                    
                p_code = doc.add_paragraph()
                p_code.paragraph_format.left_indent = Inches(0.3)
                p_code_run = p_code.add_run(code_content)
                p_code_run.font.name = 'Times New Roman'
                p_code_run.font.size = Pt(12)
                p_code.paragraph_format.space_after = Pt(10)
            else:
                # Text block (Times New Roman, 12pt)
                p_text = doc.add_paragraph()
                p_text_run = p_text.add_run(part_str)
                p_text_run.font.name = 'Times New Roman'
                p_text_run.font.size = Pt(12)
                p_text.paragraph_format.space_after = Pt(8)
                
        # 3. Validation (Times New Roman, Bold, 12pt)
        v_head = doc.add_paragraph()
        v_head_run = v_head.add_run("3. Đánh giá, chọn lọc & Tùy biến mã nguồn của Sinh viên")
        v_head_run.font.name = 'Times New Roman'
        v_head_run.font.size = Pt(12)
        v_head_run.font.bold = True
        v_head.paragraph_format.keep_with_next = True
        
        p_val = doc.add_paragraph()
        p_val_run = p_val.add_run(task["Student’s Validation / Modification"])
        p_val_run.font.name = 'Times New Roman'
        p_val_run.font.size = Pt(12)
        p_val.paragraph_format.space_after = Pt(24)
        
    os.makedirs("docs/Report/AI_Evidence", exist_ok=True)
    out_path = f"docs/Report/AI_Evidence/{week_name}.docx"
    doc.save(out_path)
    print(f"  Created Word Document: {out_path}")

if __name__ == "__main__":
    print("Converting weekly reports to Word documents in Vietnamese...")
    for week_name, tasks in WEEKLY_DATA.items():
        create_weekly_docx(week_name, tasks)
    print("All Word documents generated successfully in Vietnamese.")
